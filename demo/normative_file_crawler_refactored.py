from base_crawler import BaseCrawler
from selenium.webdriver.common.by import By
from selenium.webdriver.support import expected_conditions as EC
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import time
import os

class NormativeFileCrawler(BaseCrawler):
    """规范性文件爬虫 - 继承自BaseCrawler"""
    
    def __init__(self, download_path="./规范性文件", logger=None, task_id=None, socketio=None, progress_callback=None):
        super().__init__(download_path, logger, task_id, socketio, progress_callback)
        self.crawler_type = 'normative_file'
    
    def generate_page_url(self, base_url, page_num):
        """
        生成指定页面的URL（规范性文件只有第一页）
        :param base_url: 基础URL
        :param page_num: 页面编号（0表示第一页）
        :return: 页面URL
        """
        if page_num == 0:
            # 第一页：返回原始URL
            return base_url.rstrip('/')
        else:
            # 其他页：返回不存在的URL，让check_page_exists检测到不存在
            return f"{base_url}nonexistent_page_{page_num}.html"
    
    def get_sub_links(self, main_url):
        """
        获取主页面的子链接
        :param main_url: 主页面URL
        :return: 子链接列表
        """
        self.log(f"正在访问主页面: {main_url}")
        self.driver.get(main_url)
        time.sleep(3)
        
        sub_links = []
        try:
            self.log("尝试获取主页面链接...")
            link_elements = self.wait.until(
                EC.presence_of_all_elements_located((By.XPATH, "//div[contains(@class, 'list-right')]//ul/li/a"))
            )
            
            for element in link_elements:
                href = element.get_attribute('href')
                text = element.text.strip()
                if href and text:
                    sub_links.append({
                        'url': href,
                        'title': text
                    })
                    self.log(f"找到子链接: {text} - {href}")
            
        except Exception as e:
            self.log(f"获取子链接时出错: {e}", 'error')
        
        self.log(f"共找到 {len(sub_links)} 个子链接")
        # 更新统计信息
        self.stats['total_sub_links'] += len(sub_links)
        return sub_links
    
    def extract_text_content(self):
        """提取页面文本内容"""
        try:
            selectors = [
                "//div[@class='article-content']",
                "//div[@class='pages_content']",
                "//div[@id='content']",
                "//div[contains(@class, 'article')]"
            ]
            content_element = None
            for selector in selectors:
                try:
                    elements = self.driver.find_elements(By.XPATH, selector)
                    if elements:
                        content_element = elements[0]
                        self.log(f"使用选择器找到内容: {selector}")
                        break
                except:
                    continue
            
            if content_element:
                return content_element.text.strip()
            else:
                self.log("未找到内容区域，将提取整个body", "warning")
                return self.driver.find_element(By.TAG_NAME, 'body').text.strip()
        except Exception as e:
            self.log(f"提取文本内容时出错: {e}", "error")
            return ""
    
    def save_content_as_docx(self, title, text_content, url=None):
        """将内容保存为docx文件"""
        try:
            doc = Document()
            doc.add_heading(title, level=1).alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            
            if url:
                doc.add_paragraph(f"来源：{url}")
                doc.add_paragraph(f"抓取时间：{time.strftime('%Y-%m-%d %H:%M:%S')}")
                doc.add_paragraph("")

            if text_content:
                for para_text in text_content.split('\n'):
                    if para_text.strip():
                        doc.add_paragraph(para_text.strip())
            
            clean_title = self.clean_filename(title)
            filename = f"{clean_title}.docx"
            filepath = os.path.join(self.download_path, filename)
            
            # 如果文件已存在，直接覆盖
            if os.path.exists(filepath):
                self.log(f"文件已存在，将覆盖: {filename}")
            
            doc.save(filepath)
            self.log(f"docx文档保存成功: {filename}", "success")
            return True
        except Exception as e:
            self.log(f"保存docx文档失败: {e}", "error")
            return False
    
    def download_from_sublink(self, sub_link_info):
        """
        从子链接下载内容
        :param sub_link_info: 包含url和title的字典
        """
        url = sub_link_info['url']
        title = sub_link_info['title']
        
        self.log(f"\n正在处理: {title}")
        self.log(f"访问URL: {url}")
        
        # 统计：每处理一个子链接
        self.stats['total_documents'] += 1
        
        try:
            self.driver.get(url)
            self.wait.until(EC.presence_of_element_located((By.TAG_NAME, 'body')))
            time.sleep(2) 
            
            text_content = self.extract_text_content()
            
            if text_content:
                success = self.save_content_as_docx(title, text_content, url)
                if success:
                    self.stats['successful_downloads'] += 1
                else:
                    self.stats['failed_downloads'] += 1
                    self.stats['failed_links'].append({
                        'url': url,
                        'title': title,
                        'reason': '保存docx文档失败'
                    })
            else:
                self.log("未能提取到有效的文本内容", "warning")
                self.stats['failed_downloads'] += 1
                self.stats['failed_links'].append({
                    'url': url,
                    'title': title,
                    'reason': '未能提取到有效的文本内容'
                })

        except Exception as e:
            self.log(f"处理页面 {title} 时出错: {e}", "error")
            self.stats['failed_downloads'] += 1
            self.stats['failed_links'].append({
                'url': url,
                'title': title,
                'reason': f'处理页面时出错: {str(e)}'
            })

def main():
    """主函数"""
    print("这是一个模块文件，请通过主程序 app.py 运行。")

if __name__ == "__main__":
    main() 