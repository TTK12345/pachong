import time
import os
import re
from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from selenium.webdriver.chrome.options import Options
from urllib.parse import urljoin
import requests
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

class MemGovCrawler:
    def __init__(self, download_path=""):
        """
        初始化爬虫
        :param download_path: 下载文件保存路径
        """
        self.download_path = os.path.abspath(download_path)
        if not os.path.exists(self.download_path):
            os.makedirs(self.download_path)
        
        # 设置Chrome选项
        self.chrome_options = Options()
        self.chrome_options.add_argument('--no-sandbox')
        self.chrome_options.add_argument('--disable-dev-shm-usage')
        self.chrome_options.add_experimental_option("prefs", {
            "download.default_directory": self.download_path,
            "download.prompt_for_download": False,
            "download.directory_upgrade": True,
            "safebrowsing.enabled": True
        })
        
        self.driver = None
        self.wait = None
    
    def start_driver(self):
        """启动浏览器驱动"""
        self.driver = webdriver.Chrome(options=self.chrome_options)
        self.wait = WebDriverWait(self.driver, 10)
        print("浏览器驱动已启动")
    
    def close_driver(self):
        """关闭浏览器驱动"""
        if self.driver:
            self.driver.quit()
            print("浏览器驱动已关闭")
    
    def clean_filename(self, filename):
        """
        清理文件名，移除不合法的字符
        :param filename: 原始文件名
        :return: 清理后的文件名
        """
        # 移除或替换不合法的字符
        filename = re.sub(r'[<>:"/\\|?*]', '_', filename)
        # 移除多余的空格
        filename = re.sub(r'\s+', '_', filename.strip())
        # 限制文件名长度
        if len(filename) > 100:
            filename = filename[:100]
        return filename
    
    def get_files_in_directory(self):
        """
        获取下载目录中的所有文件
        :return: 文件列表（文件名和修改时间）
        """
        files = []
        try:
            for filename in os.listdir(self.download_path):
                if os.path.isfile(os.path.join(self.download_path, filename)):
                    file_path = os.path.join(self.download_path, filename)
                    files.append({
                        'name': filename,
                        'path': file_path,
                        'mtime': os.path.getmtime(file_path)
                    })
        except Exception as e:
            print(f"获取文件列表时出错: {e}")
        return files
    
    def wait_for_download_complete(self, initial_files, timeout=30):
        """
        等待下载完成，并返回新下载的文件
        :param initial_files: 下载前的文件列表
        :param timeout: 超时时间（秒）
        :return: 新下载的文件信息
        """
        start_time = time.time()
        print(f"开始等待下载完成，超时时间: {timeout}秒")
        
        while time.time() - start_time < timeout:
            current_files = self.get_files_in_directory()
            
            # 查找新文件
            new_files = []
            for current_file in current_files:
                is_new = True
                for initial_file in initial_files:
                    if current_file['name'] == initial_file['name']:
                        is_new = False
                        break
                if is_new:
                    new_files.append(current_file)
            
            # 检查是否有新文件且文件不再变化（下载完成）
            if new_files:
                print(f"检测到 {len(new_files)} 个新文件")
                # 等待一段时间确保文件下载完成
                time.sleep(3)
                stable_files = []
                for new_file in new_files:
                    # 检查文件是否不再变化（下载完成）
                    if not new_file['name'].endswith('.crdownload') and not new_file['name'].endswith('.tmp'):
                        stable_files.append(new_file)
                        print(f"发现稳定文件: {new_file['name']}")
                
                if stable_files:
                    return stable_files
            
            print(".", end="", flush=True)  # 显示等待进度
            time.sleep(1)
        
        print(f"\n等待超时，未检测到新文件")
        return []
    
    def rename_downloaded_file(self, file_info, new_name):
        """
        重命名下载的文件
        :param file_info: 文件信息字典
        :param new_name: 新的文件名（不包含扩展名）
        """
        try:
            old_path = file_info['path']
            old_name = file_info['name']
            
            # 获取文件扩展名
            _, ext = os.path.splitext(old_name)
            
            # 清理新文件名
            clean_name = self.clean_filename(new_name)
            
            # 构造新文件路径
            new_filename = f"{clean_name}{ext}"
            new_path = os.path.join(self.download_path, new_filename)
            
            # 如果文件已存在，添加序号
            counter = 1
            while os.path.exists(new_path):
                new_filename = f"{clean_name}_{counter}{ext}"
                new_path = os.path.join(self.download_path, new_filename)
                counter += 1
            
            # 重命名文件
            os.rename(old_path, new_path)
            print(f"文件已重命名: {old_name} -> {new_filename}")
            return new_path
            
        except Exception as e:
            print(f"重命名文件时出错: {e}")
            return file_info['path']
    
    def get_sub_links(self, main_url):
        """
        获取主页面的子链接
        :param main_url: 主页面URL
        :return: 子链接列表
        """
        print(f"正在访问主页面: {main_url}")
        self.driver.get(main_url)
        time.sleep(3)
        
        sub_links = []
        try:
            # 根据您提供的具体XPath获取所有子链接
            # 主页面链接XPath: /html/body/div[1]/div[4]/div[5]/div[1]/div[2]/ul/li[1]/a
            print("尝试获取主页面链接...")
            link_elements = self.driver.find_elements(By.XPATH, "//div[4]//div[5]//div[1]//div//ul/li/a")

            # 如果上面的XPath没有找到元素，尝试更通用的路径
            if not link_elements:
                print("尝试备用XPath路径...")
                link_elements = self.driver.find_elements(By.XPATH, "//ul/li/a[contains(@href, '.shtml')]")
                
            # 如果还是没有找到，尝试最通用的路径
            if not link_elements:
                print("尝试最通用XPath路径...")
                link_elements = self.driver.find_elements(By.XPATH, "//a[contains(@href, 'fdzdgknr') or contains(@href, 'tzgg')]")
            
            for element in link_elements:
                href = element.get_attribute('href')
                text = element.text.strip()
                if href and text:
                    sub_links.append({
                        'url': href,
                        'title': text
                    })
                    print(f"找到子链接: {text} - {href}")
            
        except Exception as e:
            print(f"获取子链接时出错: {e}")
        
        print(f"共找到 {len(sub_links)} 个子链接")
        return sub_links
    
    def extract_text_content(self, xpath_location):
        """
        从指定XPath位置提取文本内容，支持多个路径
        :param xpath_location: 主要XPath位置
        :return: 提取的文本内容和HTML结构
        """
        try:
            print(f"正在提取XPath位置的内容: {xpath_location}")
            
            # 定义多个可能的XPath选择器
            selectors = [
                xpath_location,  # 原始XPath
                "/html/body/div[4]/div[1]/div[2]/div[2]",  # 简化XPath
                "/html/body/div[1]/div[4]/div[3]/div[1]",  # 新的XPath路径1
                "/html/body/div[1]/div[4]/div[3]/div[1]/div[3]/div[1]",  # 新的XPath路径2
                "//div[@class='cont']",  # 通用内容区域
                "//div[contains(@class, 'content')]",  # 包含content类的div
                "//div[contains(@class, 'main')]",  # 主内容区域
                "//div[contains(@class, 'article')]",  # 文章内容区域
                "//main",  # main标签
                "//body"  # 整个body（最后备选）
            ]
            
            all_text_content = []
            all_html_content = []
            used_selectors = []
            
            for selector in selectors:
                try:
                    elements = self.driver.find_elements(By.XPATH, selector)
                    if elements:
                        for element in elements:
                            text_content = element.text.strip()
                            html_content = element.get_attribute('innerHTML')
                            
                            # 只添加有意义的文本内容（长度大于10个字符）
                            if text_content and len(text_content) > 10:
                                all_text_content.append(text_content)
                                all_html_content.append(html_content)
                                used_selectors.append(selector)
                                print(f"成功找到内容元素，使用选择器: {selector}")
                                print(f"提取到文本长度: {len(text_content)} 字符")
                                print(f"文本预览: {text_content[:100]}...")
                except Exception as e:
                    print(f"选择器 {selector} 提取失败: {e}")
                    continue
            
            if not all_text_content:
                print("未找到任何有效的内容")
                return None, None
            
            # 合并所有文本内容，去重并保持顺序
            combined_text = ""
            seen_content = set()
            
            for text in all_text_content:
                # 简单的去重逻辑：如果内容相似度很高，跳过
                if len(text) > 50:
                    # 对于长文本，检查是否已经包含类似内容
                    is_duplicate = False
                    for seen in seen_content:
                        if len(text) > 100 and len(seen) > 100:
                            # 如果两个长文本的前100个字符相同，认为是重复
                            if text[:100] == seen[:100]:
                                is_duplicate = True
                                break
                    
                    if not is_duplicate:
                        combined_text += text + "\n\n"
                        seen_content.add(text)
                else:
                    # 对于短文本，直接添加
                    combined_text += text + "\n\n"
                    seen_content.add(text)
            
            # 合并HTML内容（取第一个有效的）
            combined_html = all_html_content[0] if all_html_content else None
            
            print(f"最终合并文本长度: {len(combined_text)} 字符")
            print(f"使用了 {len(used_selectors)} 个选择器: {used_selectors}")
            
            return combined_text.strip(), combined_html
            
        except Exception as e:
            print(f"提取文本内容时出错: {e}")
            return None, None
    
    def save_content_as_docx(self, title, text_content, html_content=None, url=None):
        """
        将内容保存为docx文件
        :param title: 文档标题
        :param text_content: 文本内容
        :param html_content: HTML内容（可选）
        :param url: 来源URL（可选）
        :return: 是否保存成功
        """
        try:
            print("正在创建docx文档...")
            
            # 创建新的Word文档
            doc = Document()
            
            # 添加文档标题
            title_paragraph = doc.add_heading(title, level=1)
            title_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            
            # 添加来源信息
            if url:
                doc.add_paragraph(f"来源：{url}")
                doc.add_paragraph(f"抓取时间：{time.strftime('%Y-%m-%d %H:%M:%S')}")
                doc.add_paragraph("")  # 空行
            
            # 处理文本内容，按段落分割
            if text_content:
                paragraphs = text_content.split('\n')
                for para_text in paragraphs:
                    para_text = para_text.strip()
                    if para_text:  # 只添加非空段落
                        paragraph = doc.add_paragraph(para_text)
                        
                        # 如果段落看起来像标题（较短且可能包含特殊字符），设置为标题样式
                        if len(para_text) < 50 and any(char in para_text for char in ['第', '条', '章', '节', '、', '（', '）']):
                            paragraph.style = 'Heading 2'
            
            # 生成文件名
            clean_title = self.clean_filename(title)
            filename = f"{clean_title}.docx"
            filepath = os.path.join(self.download_path, filename)
            
            # 如果文件已存在，添加序号
            counter = 1
            while os.path.exists(filepath):
                filename = f"{clean_title}_{counter}.docx"
                filepath = os.path.join(self.download_path, filename)
                counter += 1
            
            # 保存文档
            doc.save(filepath)
            print(f"docx文档保存成功: {filename}")
            return True
            
        except Exception as e:
            print(f"保存docx文档失败: {e}")
            return False
    
    def download_pdf_from_sublink(self, sub_link_info):
        """
        从子链接页面提取文本内容并保存为docx，同时下载附件
        :param sub_link_info: 包含url和title的字典
        """
        url = sub_link_info['url']
        title = sub_link_info['title']
        
        print(f"\n正在处理: {title}")
        print(f"访问URL: {url}")
        
        try:
            self.driver.get(url)
            time.sleep(3)
            
            # 提取指定位置的文本内容（现在支持多个XPath路径）
            text_content, html_content = self.extract_text_content("/html/body/div[4]/div[1]/div[2]/div[2]")
            
            if text_content:
                # 保存为docx文件
                success = self.save_content_as_docx(title, text_content, html_content, url)
                if success:
                    print(f"内容已保存为docx文件: {title}")
                else:
                    print("保存docx文件失败")
            else:
                print("未能提取到有效的文本内容")
            
            # 下载页面附件
            print(f"\n开始处理页面附件...")
            attachment_count = self.download_all_attachments(title)
            if attachment_count > 0:
                print(f"成功下载 {attachment_count} 个附件")
            else:
                print("页面中没有附件或附件下载失败")
            
        except Exception as e:
            print(f"处理页面时出错: {e}")
        
        print(f"已完成处理: {title}")
    
    def check_page_exists(self, url):
        """
        检查页面是否存在
        :param url: 页面URL
        :return: 页面是否存在
        """
        try:
            response = requests.get(url, timeout=10)
            return response.status_code == 200
        except:
            return False
    
    def crawl_all_pages(self, base_url, max_pages=50):
        """
        爬取所有页面（支持翻页）
        :param base_url: 基础URL
        :param max_pages: 最大页面数，防止无限循环
        """
        page_count = 0
        
        try:
            self.start_driver()
            
            while page_count < max_pages:
                # 构造当前页面URL
                if page_count == 0:
                    current_url = base_url + "index.shtml"
                else:
                    current_url = base_url + f"index_{page_count}.shtml"
                
                print(f"\n{'='*60}")
                print(f"正在处理第 {page_count + 1} 页")
                print(f"页面URL: {current_url}")
                print(f"{'='*60}")
                
                # 检查页面是否存在
                if not self.check_page_exists(current_url):
                    print(f"第 {page_count + 1} 页不存在，停止翻页")
                    break
                
                print(f"第 {page_count + 1} 页存在，开始爬取...")
                
                # 获取当前页面的所有子链接
                sub_links = self.get_sub_links(current_url)
                
                if not sub_links:
                    print(f"第 {page_count + 1} 页未找到任何子链接")
                    page_count += 1
                    continue
                
                print(f"第 {page_count + 1} 页找到 {len(sub_links)} 个子链接")
                
                # 逐个处理子链接
                for i, sub_link in enumerate(sub_links, 1):
                    print(f"\n第 {page_count + 1} 页 - 进度: {i}/{len(sub_links)}")
                    self.download_pdf_from_sublink(sub_link)
                    time.sleep(2)  # 避免请求过于频繁
                
                print(f"第 {page_count + 1} 页处理完成")
                page_count += 1
                time.sleep(3)  # 页面间隔时间
            
            print(f"\n所有页面处理完成！共处理了 {page_count} 页")
            print(f"文件已保存到: {self.download_path}")
            
        except Exception as e:
            print(f"翻页爬取过程中出错: {e}")
        finally:
            self.close_driver()

    def crawl_all(self, main_url, max_links=None):
        """
        爬取单个页面的所有内容
        :param main_url: 主页面URL
        :param max_links: 最大处理链接数，None表示处理所有链接
        """
        try:
            if not self.driver:
                self.start_driver()
            
            # 获取所有子链接
            sub_links = self.get_sub_links(main_url)
            
            if not sub_links:
                print("未找到任何子链接")
                return
            
            # 限制处理的链接数量
            if max_links:
                sub_links = sub_links[:max_links]
                print(f"将处理前 {max_links} 个链接")
            
            # 逐个处理子链接
            for i, sub_link in enumerate(sub_links, 1):
                print(f"\n{'='*50}")
                print(f"进度: {i}/{len(sub_links)}")
                self.download_pdf_from_sublink(sub_link)
                time.sleep(2)  # 避免请求过于频繁
            
            print(f"\n当前页面任务完成！文件已保存到: {self.download_path}")
            
        except Exception as e:
            print(f"爬取过程中出错: {e}")

    def find_attachments(self, page_title=None):
        """
        在页面中查找附件链接
        :param page_title: 页面标题（用于文件命名）
        :return: 附件信息列表
        """
        attachments = []
        try:
            print("正在查找页面附件...")
            
            # 常见的附件选择器
            attachment_selectors = [
                "//a[contains(@href, '.pdf')]",  # PDF文件
                "//a[contains(@href, '.doc')]",  # Word文档
                "//a[contains(@href, '.docx')]",  # Word文档
                "//a[contains(@href, '.xls')]",  # Excel文件
                "//a[contains(@href, '.xlsx')]",  # Excel文件
                "//a[contains(@href, '.zip')]",  # 压缩文件
                "//a[contains(@href, '.rar')]",  # 压缩文件
                "//a[contains(@href, 'download')]",  # 下载链接
                "//a[contains(text(), '附件')]",  # 包含"附件"文字的链接
                "//a[contains(text(), '下载')]",  # 包含"下载"文字的链接
                "//a[contains(text(), 'PDF')]",  # 包含"PDF"文字的链接
                "//a[contains(text(), '文件')]",  # 包含"文件"文字的链接
                "//div[contains(@class, 'attachment')]//a",  # 附件区域的链接
                "//div[contains(@class, 'file')]//a",  # 文件区域的链接
                "//div[contains(@class, 'download')]//a"  # 下载区域的链接
            ]
            
            for selector in attachment_selectors:
                try:
                    elements = self.driver.find_elements(By.XPATH, selector)
                    for element in elements:
                        href = element.get_attribute('href')
                        text = element.text.strip()
                        
                        if href and text:
                            # 检查是否是有效的附件链接
                            if self.is_valid_attachment(href):
                                attachment_info = {
                                    'url': href,
                                    'text': text,
                                    'filename': self.extract_filename_from_url(href, text, page_title)
                                }
                                
                                # 避免重复添加
                                if not any(att['url'] == href for att in attachments):
                                    attachments.append(attachment_info)
                                    print(f"找到附件: {text} - {attachment_info['filename']}")
                except Exception as e:
                    print(f"选择器 {selector} 查找附件失败: {e}")
                    continue
            
            print(f"共找到 {len(attachments)} 个附件")
            return attachments
            
        except Exception as e:
            print(f"查找附件时出错: {e}")
            return []
    
    def is_valid_attachment(self, url):
        """
        检查URL是否是有效的附件链接
        :param url: 链接URL
        :return: 是否是有效附件
        """
        if not url:
            return False
        
        # 检查是否是相对路径，如果是则转换为绝对路径
        if url.startswith('/'):
            current_url = self.driver.current_url
            base_url = '/'.join(current_url.split('/')[:3])  # 获取域名部分
            url = base_url + url
        
        # 检查文件扩展名
        valid_extensions = ['.pdf', '.doc', '.docx', '.xls', '.xlsx', '.zip', '.rar', '.txt']
        if any(url.lower().endswith(ext) for ext in valid_extensions):
            return True
        
        # 检查URL是否包含下载相关关键词
        download_keywords = ['download', 'file', 'attachment', 'doc', 'pdf']
        if any(keyword in url.lower() for keyword in download_keywords):
            return True
        
        return False
    
    def extract_filename_from_url(self, url, link_text, page_title=None):
        """
        从URL或链接文本中提取文件名，包含页面标题
        :param url: 链接URL
        :param link_text: 链接文本
        :param page_title: 页面标题（可选）
        :return: 文件名
        """
        # 构建文件名前缀
        filename_parts = []
        
        # 添加页面标题（如果提供）
        if page_title:
            clean_page_title = self.clean_filename(page_title)
            filename_parts.append(clean_page_title)
        
        # 添加附件描述
        if link_text:
            # 清理链接文本，移除特殊字符
            clean_text = re.sub(r'[<>:"/\\|?*]', '_', link_text)
            clean_text = re.sub(r'\s+', '_', clean_text.strip())
            filename_parts.append(clean_text)
        
        # 如果没有链接文本，尝试从URL提取文件名
        if not link_text and '/' in url:
            url_filename = url.split('/')[-1]
            if '?' in url_filename:
                url_filename = url_filename.split('?')[0]
            if url_filename and '.' in url_filename and len(url_filename) > 0 and not url_filename.endswith('/'):
                # 移除扩展名，只保留文件名部分
                name_part = os.path.splitext(url_filename)[0]
                if name_part:
                    filename_parts.append(name_part)
        
        # 如果还是没有有效的文件名部分，使用时间戳
        if not filename_parts:
            filename_parts.append(f"附件_{int(time.time())}")
        
        # 合并文件名部分
        filename = "_".join(filename_parts)
        
        # 添加扩展名
        if '.' not in filename:
            # 检查URL中是否有文件扩展名
            if url.lower().endswith('.pdf'):
                filename += '.pdf'
            elif url.lower().endswith('.doc') or url.lower().endswith('.docx'):
                filename += '.doc'
            elif url.lower().endswith('.xls') or url.lower().endswith('.xlsx'):
                filename += '.xls'
            elif url.lower().endswith('.zip'):
                filename += '.zip'
            elif url.lower().endswith('.rar'):
                filename += '.rar'
            elif url.lower().endswith('.txt'):
                filename += '.txt'
            elif 'pdf' in link_text.lower() if link_text else False:
                filename += '.pdf'
            elif 'doc' in link_text.lower() if link_text else False:
                filename += '.doc'
            elif 'excel' in link_text.lower() or 'xls' in link_text.lower() if link_text else False:
                filename += '.xls'
            else:
                filename += '.pdf'  # 默认PDF扩展名
        
        # 限制文件名长度
        if len(filename) > 150:
            filename = filename[:150] + os.path.splitext(filename)[1]
        
        return filename
    
    def download_attachment(self, attachment_info):
        """
        下载单个附件
        :param attachment_info: 附件信息字典
        :return: 是否下载成功
        """
        url = attachment_info['url']
        filename = attachment_info['filename']
        text = attachment_info['text']
        
        try:
            print(f"正在下载附件: {text}")
            print(f"文件名: {filename}")
            print(f"下载URL: {url}")
            
            # 确保文件名有扩展名
            if '.' not in filename:
                if url.lower().endswith('.pdf'):
                    filename += '.pdf'
                elif url.lower().endswith('.doc') or url.lower().endswith('.docx'):
                    filename += '.doc'
                elif url.lower().endswith('.xls') or url.lower().endswith('.xlsx'):
                    filename += '.xls'
                elif url.lower().endswith('.zip'):
                    filename += '.zip'
                elif url.lower().endswith('.rar'):
                    filename += '.rar'
                elif url.lower().endswith('.txt'):
                    filename += '.txt'
                else:
                    filename += '.pdf'  # 默认扩展名
            
            # 获取当前页面的文件列表
            initial_files = self.get_files_in_directory()
            
            # 使用Selenium点击下载链接
            try:
                # 查找对应的链接元素
                link_elements = self.driver.find_elements(By.XPATH, f"//a[@href='{url}']")
                if link_elements:
                    link_elements[0].click()
                    print("已点击下载链接")
                else:
                    # 如果找不到精确匹配，尝试通过文本查找
                    text_elements = self.driver.find_elements(By.XPATH, f"//a[contains(text(), '{text}')]")
                    if text_elements:
                        text_elements[0].click()
                        print("已点击下载链接（通过文本匹配）")
                    else:
                        print("未找到下载链接元素")
                        return False
            except Exception as e:
                print(f"点击下载链接失败: {e}")
                return False
            
            # 等待下载完成
            new_files = self.wait_for_download_complete(initial_files, timeout=60)
            
            if new_files:
                # 重命名下载的文件
                for file_info in new_files:
                    new_path = self.rename_downloaded_file(file_info, filename)
                    print(f"附件下载完成: {os.path.basename(new_path)}")
                return True
            else:
                print("附件下载失败或超时")
                return False
                
        except Exception as e:
            print(f"下载附件时出错: {e}")
            return False
    
    def download_all_attachments(self, title):
        """
        下载页面中的所有附件
        :param title: 页面标题（用于创建附件文件夹）
        :return: 下载的附件数量
        """
        try:
            # 查找所有附件
            attachments = self.find_attachments(title)
            
            if not attachments:
                print("页面中没有找到附件")
                return 0
            
            print(f"找到 {len(attachments)} 个附件，开始下载...")
            
            # 为当前页面创建附件子文件夹
            clean_title = self.clean_filename(title)
            attachment_folder = os.path.join(self.download_path, f"{clean_title}_附件")
            if not os.path.exists(attachment_folder):
                os.makedirs(attachment_folder)
            
            print(f"附件将保存到: {attachment_folder}")
            
            # 保存当前URL，以便重新加载页面
            current_url = self.driver.current_url
            
            # 临时更改下载路径
            original_download_path = self.download_path
            self.download_path = attachment_folder
            
            # 更新Chrome下载路径
            self.chrome_options.add_experimental_option("prefs", {
                "download.default_directory": attachment_folder,
                "download.prompt_for_download": False,
                "download.directory_upgrade": True,
                "safebrowsing.enabled": True
            })
            
            # 重新启动驱动以应用新的下载路径
            if self.driver:
                self.close_driver()
                self.start_driver()
                self.driver.get(current_url)  # 重新加载当前页面
                time.sleep(3)
            
            downloaded_count = 0
            
            # 逐个下载附件
            for i, attachment in enumerate(attachments, 1):
                print(f"\n下载进度: {i}/{len(attachments)}")
                if self.download_attachment(attachment):
                    downloaded_count += 1
                time.sleep(2)  # 避免下载过于频繁
            
            # 恢复原始下载路径
            self.download_path = original_download_path
            self.chrome_options.add_experimental_option("prefs", {
                "download.default_directory": original_download_path,
                "download.prompt_for_download": False,
                "download.directory_upgrade": True,
                "safebrowsing.enabled": True
            })
            
            # 重新启动驱动以恢复原始下载路径
            if self.driver:
                self.close_driver()
                self.start_driver()
                self.driver.get(current_url)  # 重新加载当前页面
                time.sleep(3)
            
            print(f"附件下载完成！成功下载 {downloaded_count}/{len(attachments)} 个附件")
            print(f"附件保存在: {attachment_folder}")
            
            return downloaded_count
            
        except Exception as e:
            print(f"下载附件过程中出错: {e}")
            return 0

def main():
    """主函数"""
    main_url = "https://www.mem.gov.cn/fw/flfgbz/gfxwj/"
    
    # 创建爬虫实例
    crawler = MemGovCrawler(download_path="./规范性文件")
    
    # 直接开始爬取所有页面（自动翻页）
    print("开始爬取所有页面并保存为docx格式...")
    max_pages = 50  # 默认最大页面数
    
    try:
        crawler.crawl_all_pages(main_url, max_pages=max_pages)
    except KeyboardInterrupt:
        print("\n用户中断爬取")
    except Exception as e:
        print(f"\n爬取过程中出错: {e}")

if __name__ == "__main__":
    main() 