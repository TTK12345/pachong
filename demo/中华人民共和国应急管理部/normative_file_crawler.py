import time
import os
import re
from selenium.webdriver.common.by import By
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from selenium.webdriver.chrome.options import Options
from urllib.parse import urljoin
import requests
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

class NormativeFileCrawler:
    def __init__(self, download_path, logger=None):
        self.download_path = os.path.abspath(download_path)
        if not os.path.exists(self.download_path):
            os.makedirs(self.download_path)
        
        self.logger = logger
        self.driver = None
        self.wait = None

    def log(self, message, level='info'):
        if self.logger:
            self.logger.log(message, level)
        else:
            print(f"[{level.upper()}] {message}")
    
    def clean_filename(self, filename):
        filename = re.sub(r'[\\/*?:"<>|]', '', filename)
        filename = re.sub(r'\s+', '_', filename.strip())
        return filename[:100]

    def get_sub_links(self, main_url):
        self.log(f"正在访问主页面: {main_url}")
        self.driver.get(main_url)
        time.sleep(3)
        
        sub_links = []
        try:
            self.log("尝试获取主页面链接...")
            link_elements = self.wait.until(
                EC.presence_of_all_elements_located((By.XPATH, "//div[contains(@class, 'list-right')]//ul/li/a"))
            )
            
            for element in link_elements:
                href = element.get_attribute('href')
                text = element.text.strip()
                if href and text:
                    sub_links.append({'url': href, 'title': text})
            
        except Exception as e:
            self.log(f"获取子链接时出错: {e}", 'error')
        
        self.log(f"共找到 {len(sub_links)} 个子链接")
        return sub_links
    
    def extract_text_content(self):
        try:
            selectors = [
                "//div[@class='article-content']",
                "//div[@class='pages_content']",
                "//div[@id='content']",
                "//div[contains(@class, 'article')]"
            ]
            content_element = None
            for selector in selectors:
                try:
                    elements = self.driver.find_elements(By.XPATH, selector)
                    if elements:
                        content_element = elements[0]
                        self.log(f"使用选择器找到内容: {selector}")
                        break
                except:
                    continue
            
            if content_element:
                return content_element.text.strip()
            else:
                self.log("未找到内容区域，将提取整个body", "warning")
                return self.driver.find_element(By.TAG_NAME, 'body').text.strip()
        except Exception as e:
            self.log(f"提取文本内容时出错: {e}", "error")
            return ""

    def save_content_as_docx(self, title, text_content, url=None):
        try:
            doc = Document()
            doc.add_heading(title, level=1).alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            
            if url:
                doc.add_paragraph(f"来源：{url}")
                doc.add_paragraph(f"抓取时间：{time.strftime('%Y-%m-%d %H:%M:%S')}")
                doc.add_paragraph("")

            if text_content:
                for para_text in text_content.split('\n'):
                    if para_text.strip():
                        doc.add_paragraph(para_text.strip())
            
            clean_title = self.clean_filename(title)
            filename = f"{clean_title}.docx"
            filepath = os.path.join(self.download_path, filename)
            
            counter = 1
            while os.path.exists(filepath):
                filename = f"{clean_title}_{counter}.docx"
                filepath = os.path.join(self.download_path, filename)
                counter += 1
            
            doc.save(filepath)
            self.log(f"docx文档保存成功: {filename}", "success")
            return True
        except Exception as e:
            self.log(f"保存docx文档失败: {e}", "error")
            return False

    def download_pdf_from_sublink(self, sub_link_info):
        url = sub_link_info['url']
        title = sub_link_info['title']
        
        try:
            self.driver.get(url)
            self.wait.until(EC.presence_of_element_located((By.TAG_NAME, 'body')))
            time.sleep(2) 
            
            text_content = self.extract_text_content()
            
            if text_content:
                self.save_content_as_docx(title, text_content, url)
            else:
                self.log("未能提取到有效的文本内容", "warning")

        except Exception as e:
            self.log(f"处理页面 {title} 时出错: {e}", "error")

def main():
    print("这是一个模块文件，请通过主程序 app.py 运行。")

if __name__ == "__main__":
    main() 